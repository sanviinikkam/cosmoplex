"""
Admin portal endpoints (hidden /admin). Single shared password → admin JWT.

Everything except /admin/login requires a valid admin token (require_admin).

  POST   /admin/login                         — password → admin token
  GET    /admin/courses                       — full course tree
  POST   /admin/courses                       — create course
  PUT    /admin/courses/{id}                  — rename / describe
  DELETE /admin/courses/{id}                  — delete course (cascades)
  POST   /admin/modules                       — add module to a course
  PUT    /admin/modules/{id} / DELETE
  POST   /admin/sections                      — add section to a module
  PUT    /admin/sections/{id} / DELETE
  POST   /admin/videos                        — add lesson to a section
  PUT    /admin/videos/{id} / DELETE
  PUT    /admin/videos/{id}/variant           — set a per-language Cloudinary id
  DELETE /admin/videos/{id}/variant/{lang}
  POST   /admin/cloudinary/signature          — signed direct-upload params
"""
import asyncio
import hashlib
import hmac
import httpx
import io
import json
import re
import time
from collections import defaultdict, deque
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Request, UploadFile, File, Form
from pydantic import BaseModel
from sqlalchemy import delete, func, or_, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from core.auth import create_admin_token, require_admin
from core.config import settings
from db.database import get_db
from api.whatsapp_drip import SIGNUP_STAGES
from db.models import (
    Course, CourseModule, Section, Video, VideoLanguageVariant, VideoProgress,
    QuizQuestion, AssignmentPrompt, IntroVideo, MarketingAsset,
    LearnerProfile, WhatsAppSession, WhatsAppMessage, Certificate,
    ExamAttempt, LessonAssignmentSubmission, ModuleProgress, Referral,
)

router = APIRouter(prefix="/admin", tags=["admin"])

SUPPORTED_LANGUAGES = {"en", "hi", "mr", "te", "ta", "kn"}

# Cloudinary IDs that already exist and are used by the learner site / WhatsApp,
# keyed by lesson title. /admin/sync-videos imports these into the DB so they
# show up (and become editable) in the portal.
KNOWN_VIDEO_IDS = {
    "The 10 AI Words Every Fresher Must Know": {
        "en": "2.1_English_compressed_s6vhdd",
        "hi": "2.1_hindi_sixgnf",
        "mr": "2.1_Marathi_cws5fc",
        "te": "2.1_Telugu_qloes6",
        "ta": "2.1_tamil_tl4rf2",
        "kn": "2.1_Kannada_azgabe",
    },
    "When AI Confidently Lies - Hallucination": {
        "hi": "2.4_hindi_compressed_vxkloy",
    },
}


# ── Auth ──────────────────────────────────────────────────────────────────────
class LoginBody(BaseModel):
    password: str


# ── Brute-force throttle for /admin/login ─────────────────────────────────────
# In-memory, per-process (same caveat as core/rate_limit.py — fine for the single
# Render instance; move to Redis if this ever scales horizontally). After
# _LOGIN_MAX_FAILS failed attempts from one IP within a sliding window, further
# attempts are rejected with 429 until the oldest failure ages out.
_LOGIN_WINDOW = 900        # 15-minute sliding window
_LOGIN_MAX_FAILS = 5       # lock the IP after this many failures in the window
_login_fails: dict[str, deque] = defaultdict(deque)


def _client_ip(request: Request) -> str:
    # Behind Render/Vercel the real client IP is the first X-Forwarded-For entry.
    xff = request.headers.get("x-forwarded-for")
    if xff:
        return xff.split(",")[0].strip()
    return request.client.host if request.client else "unknown"


def _login_locked(ip: str, now: float) -> bool:
    dq = _login_fails[ip]
    while dq and now - dq[0] > _LOGIN_WINDOW:
        dq.popleft()
    if not dq:
        _login_fails.pop(ip, None)
        return False
    return len(dq) >= _LOGIN_MAX_FAILS


@router.post("/login")
async def admin_login(body: LoginBody, request: Request):
    ip = _client_ip(request)
    now = time.time()
    if _login_locked(ip, now):
        raise HTTPException(
            status_code=429,
            detail="Too many failed attempts. Please wait a few minutes and try again.",
            headers={"Retry-After": str(_LOGIN_WINDOW)},
        )
    # Constant-time compare so response timing can't leak the password.
    ok = bool(settings.admin_password) and hmac.compare_digest(
        body.password.encode("utf-8"), settings.admin_password.encode("utf-8"))
    if not ok:
        _login_fails[ip].append(now)
        raise HTTPException(status_code=401, detail="Incorrect password")
    _login_fails.pop(ip, None)   # reset the counter on a successful login
    return {"access_token": create_admin_token(), "token_type": "bearer"}


# ── Serialization ──────────────────────────────────────────────────────────────
def _video_dict(v: Video) -> dict:
    return {
        "id": v.id,
        "title": v.title,
        "orderIndex": v.order_index,
        "baseCloudinaryId": v.cloudinary_public_id,
        "durationSeconds": v.duration_seconds,
        "variants": [
            {"language": lv.language, "cloudinaryPublicId": lv.cloudinary_public_id,
             "durationSeconds": lv.duration_seconds}
            for lv in sorted(v.language_variants, key=lambda x: x.language)
        ],
    }


def _course_tree(c: Course) -> dict:
    return {
        "id": c.id,
        "title": c.title,
        "description": c.description,
        "thumbnailCloudinaryId": c.thumbnail_cloudinary_id,
        "modules": [
            {
                "id": m.id, "title": m.title, "outcome": m.outcome,
                "orderIndex": m.order_index, "level": m.level,
                "hasContentDoc": bool(m.content_doc),
                "contentDocPreview": (m.content_doc or "")[:200],
                "sections": [
                    {
                        "id": s.id, "title": s.title, "orderIndex": s.order_index,
                        "videos": [_video_dict(v) for v in sorted(s.videos, key=lambda x: x.order_index)],
                    }
                    for s in sorted(m.sections, key=lambda x: x.order_index)
                ],
            }
            for m in sorted(c.modules, key=lambda x: x.order_index)
        ],
    }


_FULL_TREE = (
    selectinload(Course.modules)
    .selectinload(CourseModule.sections)
    .selectinload(Section.videos)
    .selectinload(Video.language_variants)
)


async def _load_course(course_id: str, db: AsyncSession) -> Course:
    res = await db.execute(select(Course).options(_FULL_TREE).where(Course.id == course_id))
    course = res.scalar_one_or_none()
    if not course:
        raise HTTPException(status_code=404, detail="Course not found")
    return course


async def _next_order(db: AsyncSession, model, fk_col, fk_val) -> int:
    res = await db.execute(select(func.coalesce(func.max(model.order_index), -1)).where(fk_col == fk_val))
    return int(res.scalar() or -1) + 1


# ── Courses ─────────────────────────────────────────────────────────────────────
@router.get("/dashboard")
async def dashboard(_: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    """System overview for the admin portal: content counts, web-learner and
    WhatsApp-user stats, status breakdowns, and recent activity. Defensive —
    a failure in one block returns {"error": ...} for that block only."""
    from datetime import datetime, timedelta
    now = datetime.utcnow()

    async def count(model, *where):
        q = select(func.count()).select_from(model)
        for w in where:
            q = q.where(w)
        return (await db.execute(q)).scalar() or 0

    async def group(col):
        rows = (await db.execute(select(col, func.count()).group_by(col))).all()
        return {(str(k) if k not in (None, "") else "—"): n for k, n in rows}

    # ── Content ──
    try:
        content = {
            "courses": await count(Course), "modules": await count(CourseModule),
            "sections": await count(Section), "videos": await count(Video),
            "quizzes": await count(QuizQuestion), "assignments": await count(AssignmentPrompt),
        }
    except Exception as e:
        content = {"error": type(e).__name__}

    # ── Web learners ──
    try:
        with_progress = (await db.execute(
            select(func.count(func.distinct(VideoProgress.learner_id))))).scalar() or 0
        recent = (await db.execute(
            select(LearnerProfile).order_by(LearnerProfile.created_at.desc()).limit(8))).scalars().all()
        web = {
            "total": await count(LearnerProfile),
            "testAccounts": await count(LearnerProfile, LearnerProfile.is_test.is_(True)),
            "certificates": await count(Certificate),
            "withProgress": with_progress,
            "byLanguage": await group(LearnerProfile.preferred_language),
            "recent": [{
                "id": r.id,
                "name": r.name, "email": r.email, "language": r.preferred_language,
                "certificate": bool(r.certificate_issued), "isTest": bool(r.is_test),
                "score": r.total_score,
                "joined": r.created_at.isoformat() if r.created_at else None,
            } for r in recent],
        }
    except Exception as e:
        web = {"error": type(e).__name__}

    # ── WhatsApp users ──
    try:
        recent_wa = (await db.execute(
            select(WhatsAppSession).order_by(WhatsAppSession.last_active_at.desc()).limit(10))).scalars().all()
        _dash_labels = await _lesson_labels(db)
        def mask(p):
            return ("•••• " + p[-4:]) if p and len(p) >= 4 else (p or "—")
        whatsapp = {
            "total": await count(WhatsAppSession),
            "active24h": await count(WhatsAppSession, WhatsAppSession.last_active_at >= now - timedelta(hours=24)),
            "active7d": await count(WhatsAppSession, WhatsAppSession.last_active_at >= now - timedelta(days=7)),
            "completed": await count(WhatsAppSession, WhatsAppSession.stage == "done"),
            "byStage": await group(WhatsAppSession.stage),
            "byLanguage": await group(WhatsAppSession.language),
            # Same microlesson label the user directory shows. This payload shares
            # the WaSessionRow type with /admin/users, so sending a raw index here
            # would make that type a lie for anything that later reads the field.
            "recent": [{
                "id": r.phone,
                "name": r.name or "—", "phone": mask(r.phone), "language": r.language,
                "stage": r.stage,
                "lesson": (_dash_labels[r.lesson_index or 0]
                           if 0 <= (r.lesson_index or 0) < len(_dash_labels) else None),
                "lessonIndex": r.lesson_index,
                "lastActive": r.last_active_at.isoformat() if r.last_active_at else None,
            } for r in recent_wa],
        }
    except Exception as e:
        whatsapp = {"error": type(e).__name__}

    return {"generatedAt": now.isoformat(), "content": content, "web": web, "whatsapp": whatsapp}


@router.get("/system-check")
async def system_check(request: Request, _: bool = Depends(require_admin),
                       db: AsyncSession = Depends(get_db)):
    """One-shot health check across every dependency. Each entry is
    {key, label, status: ok|warn|error, detail}. `overall` = worst status.
    External calls are cheap/free (validation pings) and time-boxed."""
    from datetime import datetime
    from sqlalchemy import text as _text
    checks: list[dict] = []

    def rank(s): return {"ok": 0, "warn": 1, "error": 2}.get(s, 2)

    # 1. Database
    try:
        await db.execute(_text("SELECT 1"))
        c = {m.__name__: (await db.execute(select(func.count()).select_from(m))).scalar() or 0
             for m in (Course, Video, QuizQuestion, AssignmentPrompt)}
        checks.append({"key": "database", "label": "Database (Neon)", "status": "ok",
                       "detail": f"connected · {c['Course']} courses, {c['Video']} videos, "
                                 f"{c['QuizQuestion']} quizzes, {c['AssignmentPrompt']} assignments"})
    except Exception as e:
        checks.append({"key": "database", "label": "Database (Neon)", "status": "error",
                       "detail": f"cannot connect ({type(e).__name__})"})

    # 2. Anthropic (Claude) — free key-validation ping via the models endpoint
    if not settings.anthropic_api_key:
        checks.append({"key": "anthropic", "label": "Claude (Anthropic)", "status": "warn", "detail": "API key not set"})
    else:
        try:
            async with httpx.AsyncClient(timeout=10) as h:
                r = await h.get("https://api.anthropic.com/v1/models",
                                headers={"x-api-key": settings.anthropic_api_key,
                                         "anthropic-version": "2023-06-01"})
            ok = r.status_code < 400
            checks.append({"key": "anthropic", "label": "Claude (Anthropic)", "status": "ok" if ok else "error",
                           "detail": "key valid, API reachable" if ok else f"key rejected (HTTP {r.status_code})"})
        except Exception as e:
            checks.append({"key": "anthropic", "label": "Claude (Anthropic)", "status": "error",
                           "detail": f"unreachable ({type(e).__name__})"})

    # 3. Groq (voice → text) — free models list
    if not settings.groq_api_key:
        checks.append({"key": "groq", "label": "Groq (voice)", "status": "warn", "detail": "API key not set (voice notes disabled)"})
    else:
        try:
            async with httpx.AsyncClient(timeout=10) as h:
                r = await h.get("https://api.groq.com/openai/v1/models",
                                headers={"Authorization": f"Bearer {settings.groq_api_key}"})
            checks.append({"key": "groq", "label": "Groq (voice)", "status": "ok" if r.status_code < 400 else "error",
                           "detail": "key valid, API reachable" if r.status_code < 400 else f"HTTP {r.status_code}"})
        except Exception as e:
            checks.append({"key": "groq", "label": "Groq (voice)", "status": "error", "detail": f"unreachable ({type(e).__name__})"})

    # 4. WhatsApp (Meta Cloud API) — validate the token against the phone-number node
    if not (settings.whatsapp_token and settings.whatsapp_phone_number_id):
        checks.append({"key": "whatsapp", "label": "WhatsApp (Meta)", "status": "warn", "detail": "token / phone number id not set"})
    else:
        try:
            url = f"https://graph.facebook.com/{settings.graph_api_version}/{settings.whatsapp_phone_number_id}"
            async with httpx.AsyncClient(timeout=10) as h:
                r = await h.get(url, params={"fields": "verified_name,quality_rating,messaging_limit_tier"},
                                headers={"Authorization": f"Bearer {settings.whatsapp_token}"})
            if r.status_code < 400:
                j = r.json()
                tier = (j.get("messaging_limit_tier") or "?").replace("TIER_", "")
                checks.append({"key": "whatsapp", "label": "WhatsApp (Meta)", "status": "ok",
                               "detail": f"token valid · {j.get('verified_name', 'number')} · quality: {j.get('quality_rating', '?')} · daily tier: {tier}"})
            else:
                checks.append({"key": "whatsapp", "label": "WhatsApp (Meta)", "status": "error",
                               "detail": f"token rejected (HTTP {r.status_code}) — the bot cannot send/receive until this is fixed"})
        except Exception as e:
            checks.append({"key": "whatsapp", "label": "WhatsApp (Meta)", "status": "error", "detail": f"unreachable ({type(e).__name__})"})

    # 5. Cloudinary — real usage / quota via the Admin API
    if not (settings.cloudinary_cloud_name and settings.cloudinary_api_key and settings.cloudinary_api_secret):
        checks.append({"key": "cloudinary", "label": "Cloudinary (media)", "status": "warn", "detail": "credentials incomplete"})
    else:
        try:
            async with httpx.AsyncClient(timeout=10) as h:
                r = await h.get(f"https://api.cloudinary.com/v1_1/{settings.cloudinary_cloud_name}/usage",
                                auth=(settings.cloudinary_api_key, settings.cloudinary_api_secret))
            if r.status_code < 400:
                j = r.json()
                plan = j.get("plan", "?")
                pct = ((j.get("credits") or {}).get("used_percent"))
                if pct is None:
                    status, detail = "ok", f"reachable · plan '{plan}'"
                else:
                    status = "error" if pct >= 95 else "warn" if pct >= 80 else "ok"
                    detail = f"{pct:.0f}% of monthly quota used · plan '{plan}'"
                checks.append({"key": "cloudinary", "label": "Cloudinary (media)", "status": status, "detail": detail})
            else:
                checks.append({"key": "cloudinary", "label": "Cloudinary (media)", "status": "warn",
                               "detail": f"configured, but usage API returned HTTP {r.status_code}"})
        except Exception as e:
            checks.append({"key": "cloudinary", "label": "Cloudinary (media)", "status": "warn",
                           "detail": f"configured · usage check failed ({type(e).__name__})"})

    # 6. Webhook signature verification
    if settings.whatsapp_app_secret:
        checks.append({"key": "webhook_security", "label": "Webhook signature", "status": "ok", "detail": "verification ON"})
    else:
        checks.append({"key": "webhook_security", "label": "Webhook signature", "status": "warn",
                       "detail": "WHATSAPP_APP_SECRET not set — inbound webhook signature check is OFF"})

    # 6b. Ops-endpoint key. Without it the ops endpoints (/run-drip, /setup,
    # /register, /subscribe, /diag*) refuse everything — including an external
    # cron job — so surface it here rather than letting it fail silently.
    if settings.whatsapp_ops_key:
        checks.append({"key": "ops_key", "label": "Ops endpoints", "status": "ok",
                       "detail": "protected by WHATSAPP_OPS_KEY"})
    else:
        checks.append({"key": "ops_key", "label": "Ops endpoints", "status": "warn",
                       "detail": "WHATSAPP_OPS_KEY not set — /run-drip, /setup, /register, "
                                 "/subscribe and /diag* refuse every request"})

    # 7. Drip / nudge scheduler
    sched = getattr(request.app.state, "scheduler", None)
    running = bool(sched and getattr(sched, "running", False))
    checks.append({"key": "scheduler", "label": "Nudge scheduler", "status": "ok" if running else "warn",
                   "detail": "hourly drip running" if running else "not running (nudges won't auto-fire in-process)"})

    # 8. Recent AI call failures — the reactive signal for Anthropic/Groq credit/quota
    # problems (there's no balance API to read proactively).
    try:
        from core.ai_health import recent_errors
        errs = recent_errors()
        if not errs:
            checks.append({"key": "ai_errors", "label": "AI call failures", "status": "ok",
                           "detail": "no AI errors in the last hour"})
        else:
            billing = any(v["billingLikely"] for v in errs.values())
            parts = [f"{p}: {v['error'][:70]} ({v['minutesAgo']}m ago)" for p, v in errs.items()]
            checks.append({"key": "ai_errors", "label": "AI call failures",
                           "status": "error" if billing else "warn",
                           "detail": ("⚠ looks credit/billing/quota-related — " if billing else "") + "; ".join(parts)})
    except Exception:
        pass

    # 9. AI spend budget (today)
    try:
        from core import spend_guard
        used = spend_guard._count if spend_guard._day_key == spend_guard._today() else 0
        limit = settings.daily_ai_call_limit
        checks.append({"key": "ai_budget", "label": "AI spend budget", "status": "warn" if used >= limit else "ok",
                       "detail": f"{used}/{limit} AI calls used today" + (" — cap reached" if used >= limit else "")})
    except Exception:
        pass

    overall = "ok"
    for c in checks:
        if rank(c["status"]) > rank(overall):
            overall = c["status"]
    return {"generatedAt": datetime.utcnow().isoformat(), "environment": settings.environment,
            "overall": overall, "checks": checks}


async def _lesson_labels(db) -> list[str]:
    """Microlesson labels ("1.1", "1.2", "2.1", ...) in course order.

    Reuses the SAME label _db_lessons builds for the learner-facing flow
    (module.order_index + 1 . section.order_index + 1), so the admin table can
    never drift into a second numbering scheme. Labels come from module/section
    order, which is language-independent, so 'en' is safe to look them up with.
    """
    from api.whatsapp_routes import _db_lessons   # lazy: avoids a circular import
    try:
        lessons = await _db_lessons(db, "en")
    except Exception:
        return []
    return [(l.get("label") or "") for l in lessons]


async def _lesson_facet(db) -> list[str]:
    """Microlesson labels that at least one learner is currently on."""
    labels = await _lesson_labels(db)
    idxs = (await db.execute(select(WhatsAppSession.lesson_index).distinct()
                            .where(WhatsAppSession.lesson_index.is_not(None)))).scalars().all()
    present = sorted({int(i) for i in idxs})
    return [labels[i] for i in present if 0 <= i < len(labels) and labels[i]]


async def _facets(db, channel: str) -> dict:
    """Distinct values actually present, per filterable column.

    The column dropdowns are built from this, not from a hardcoded list: a
    hardcoded list shows stages nobody is in (and silently misses any new one),
    which makes the filter feel broken when a pick returns zero rows.

    Computed over the WHOLE channel, deliberately not narrowed by the other
    filters — otherwise choosing a stage would empty the language dropdown and
    the learner could not switch without clearing everything first.
    """
    async def distinct(col):
        rows = (await db.execute(select(col).distinct().where(col.is_not(None)))).scalars().all()
        vals = [v for v in rows if v is not None and str(v).strip() != ""]
        return sorted(vals, key=lambda v: (str(v).lower() if isinstance(v, str) else v))

    if channel == "whatsapp":
        return {
            "stage": await distinct(WhatsAppSession.stage),
            "language": await distinct(WhatsAppSession.language),
            "source_type": await distinct(WhatsAppSession.source_type),
            "campaign": await distinct(WhatsAppSession.campaign),
            # Offer the microlesson labels learners actually sit on, in COURSE
            # order. Sorting the strings would put "1.10" before "1.2".
            "lesson": await _lesson_facet(db),
        }
    return {
        "language": await distinct(LearnerProfile.preferred_language),
        "certificate": ["yes", "no"],   # a boolean column: both states are meaningful
    }


def _date_range(from_date: str | None, to_date: str | None):
    """Parse YYYY-MM-DD bounds into datetimes. `to` is inclusive of that whole
    day, so a single-day filter (from == to) returns that day rather than nothing."""
    from datetime import datetime, timedelta
    start = end = None
    for raw, is_start in ((from_date, True), (to_date, False)):
        if not raw or not raw.strip():
            continue
        try:
            d = datetime.strptime(raw.strip()[:10], "%Y-%m-%d")
        except ValueError:
            continue          # ignore junk rather than 500 the admin page
        if is_start:
            start = d
        else:
            end = d + timedelta(days=1)
    return start, end


@router.get("/campaigns")
async def campaign_report(
    from_date: str | None = None,
    to_date: str | None = None,
    _: bool = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    """Acquisition funnel per campaign.

    Counts alone can't tell you which campaign is worth more money — a campaign
    that sends 500 clicks who never answer is worse than one that sends 50 who
    finish. So each row carries the whole funnel, not just arrivals.
    """
    # Filter on when the learner ARRIVED. first_seen_at only exists on rows created
    # after attribution shipped, so fall back to created_at — otherwise every older
    # learner would silently vanish from a date-filtered report.
    start, end = _date_range(from_date, to_date)
    arrived = func.coalesce(WhatsAppSession.first_seen_at, WhatsAppSession.created_at)
    base = select(WhatsAppSession)
    if start is not None:
        base = base.where(arrived >= start)
    if end is not None:
        base = base.where(arrived < end)
    rows = (await db.execute(base)).scalars().all()
    buckets: dict[str, dict] = {}
    for r in rows:
        key = r.campaign or "organic"
        b = buckets.setdefault(key, {
            "campaign": key,
            "source_type": r.source_type or "organic",
            "headline": r.source_headline,
            "ad_id": r.ad_id,
            "arrived": 0, "picked_language": 0, "signed_up": 0,
            "started_lesson": 0, "completed": 0, "opted_out": 0,
        })
        b["arrived"] += 1
        if r.language:
            b["picked_language"] += 1
        # "Signed up" = got THROUGH onboarding. It must not key off r.name:
        # WhatsApp sends the sender's profile name in the webhook and we store it
        # on their very first message, so name is set for everyone who ever writes
        # in — which made this column read 100% on every campaign and told the
        # operator nothing.
        if r.stage not in SIGNUP_STAGES:
            b["signed_up"] += 1
        if (r.lesson_index or 0) > 0 or r.stage in ("lesson", "quiz", "quiz_failed",
                                                    "assignment", "between_lessons",
                                                    "clarify", "done"):
            b["started_lesson"] += 1
        if r.stage == "done":
            b["completed"] += 1
        if getattr(r, "opt_out", False):
            b["opted_out"] += 1
        # Keep the most descriptive label seen for this campaign.
        if r.source_headline and not b["headline"]:
            b["headline"] = r.source_headline
        if r.ad_id and not b["ad_id"]:
            b["ad_id"] = r.ad_id

    out = sorted(buckets.values(), key=lambda b: b["arrived"], reverse=True)
    for b in out:
        a = b["arrived"] or 1
        b["signup_rate"] = round(100 * b["signed_up"] / a)
        b["completion_rate"] = round(100 * b["completed"] / a)
    return {"campaigns": out, "total_users": len(rows),
            "from_date": from_date or None, "to_date": to_date or None}


@router.get("/users")
async def all_users(
    channel: str = "web",
    q: str | None = None,
    # Per-column filters. Applied server-side on purpose: filtering only the
    # loaded page would show "3 results" out of a 500-row page and read as the
    # whole truth.
    from_date: str | None = None,   # joined >= (YYYY-MM-DD)
    to_date: str | None = None,     # joined <= (inclusive day)
    language: str | None = None,
    stage: str | None = None,
    campaign: str | None = None,
    source_type: str | None = None,
    lesson: str | None = None,          # microlesson label, e.g. "1.3"
    active_within_days: int | None = None,
    certificate: str | None = None,   # web only: "yes" | "no"
    limit: int = 500,
    offset: int = 0,
    _: bool = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    """The FULL user directory (not just recent) for the admin portal, one channel
    at a time. channel=web|whatsapp, optional case-insensitive search `q`, paginated.
    Rows carry the same `id` the detail endpoints (/learner/{id}, /whatsapp/{phone})
    expect, so a row stays click-through to the per-user modal."""
    limit = max(1, min(limit, 1000))
    term = f"%{(q or '').strip().lower()}%"

    if channel == "whatsapp":
        def mask(p):
            return ("•••• " + p[-4:]) if p and len(p) >= 4 else (p or "—")
        base = select(WhatsAppSession)
        if q and q.strip():
            base = base.where(or_(
                func.lower(WhatsAppSession.name).like(term),
                WhatsAppSession.phone.like(term),
            ))
        start, end = _date_range(from_date, to_date)
        if start is not None:
            base = base.where(WhatsAppSession.created_at >= start)
        if end is not None:
            base = base.where(WhatsAppSession.created_at < end)
        if language and language.strip():
            base = base.where(WhatsAppSession.language == language.strip())
        if stage and stage.strip():
            base = base.where(WhatsAppSession.stage == stage.strip())
        if source_type and source_type.strip():
            base = base.where(WhatsAppSession.source_type == source_type.strip())
        if campaign and campaign.strip():
            base = base.where(WhatsAppSession.campaign.ilike(f"%{campaign.strip()}%"))
        if lesson and lesson.strip():
            # The dropdown sends a microlesson label ("1.3"); resolve it to the
            # 0-based index the column actually stores.
            labels_for_filter = await _lesson_labels(db)
            want = lesson.strip()
            try:
                base = base.where(WhatsAppSession.lesson_index == labels_for_filter.index(want))
            except ValueError:
                # Unknown label — match nothing rather than silently ignoring it,
                # so a stale bookmark cannot look like "no filter applied".
                base = base.where(WhatsAppSession.lesson_index == -1)
        if active_within_days is not None and active_within_days > 0:
            from datetime import datetime, timedelta
            cutoff = datetime.utcnow() - timedelta(days=active_within_days)
            base = base.where(WhatsAppSession.last_active_at >= cutoff)
        total = (await db.execute(
            select(func.count()).select_from(base.subquery()))).scalar() or 0
        rows = (await db.execute(
            base.order_by(WhatsAppSession.last_active_at.desc().nullslast())
                .offset(max(0, offset)).limit(limit))).scalars().all()
        labels = await _lesson_labels(db)

        def lesson_label(i):
            i = i or 0
            return labels[i] if 0 <= i < len(labels) and labels[i] else str(i + 1)

        items = [{
            "id": r.phone,
            "name": r.name or "—", "phone": mask(r.phone), "language": r.language,
            "stage": r.stage, "lesson": lesson_label(r.lesson_index),
            "lessonIndex": r.lesson_index,
            "campaign": r.campaign, "sourceType": r.source_type,
            "lastActive": r.last_active_at.isoformat() if r.last_active_at else None,
            "joined": r.created_at.isoformat() if getattr(r, "created_at", None) else None,
        } for r in rows]
    else:
        base = select(LearnerProfile)
        if q and q.strip():
            base = base.where(or_(
                func.lower(LearnerProfile.name).like(term),
                func.lower(LearnerProfile.email).like(term),
            ))
        start, end = _date_range(from_date, to_date)
        if start is not None:
            base = base.where(LearnerProfile.created_at >= start)
        if end is not None:
            base = base.where(LearnerProfile.created_at < end)
        if language and language.strip():
            base = base.where(LearnerProfile.preferred_language == language.strip())
        if certificate in ("yes", "no"):
            base = base.where(LearnerProfile.certificate_issued.is_(certificate == "yes"))
        total = (await db.execute(
            select(func.count()).select_from(base.subquery()))).scalar() or 0
        rows = (await db.execute(
            base.order_by(LearnerProfile.created_at.desc().nullslast())
                .offset(max(0, offset)).limit(limit))).scalars().all()
        items = [{
            "id": r.id,
            "name": r.name, "email": r.email, "language": r.preferred_language,
            "certificate": bool(r.certificate_issued), "isTest": bool(r.is_test),
            "score": r.total_score,
            "joined": r.created_at.isoformat() if r.created_at else None,
        } for r in rows]

    return {
        "channel": channel, "total": total,
        "facets": await _facets(db, channel),
        "offset": max(0, offset), "limit": limit,
        "count": len(items), "items": items,
    }


@router.get("/referrals")
async def list_referrals(_: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    """Referral ledger for the admin portal: totals + recent rows with the
    referrer's contact (email / masked phone) for payout."""
    rows = (await db.execute(select(Referral).order_by(Referral.created_at.desc()).limit(200))).scalars().all()
    items = []
    for r in rows:
        if r.referrer_kind == "web":
            lp = await db.get(LearnerProfile, r.referrer_id)
            contact = (lp.email if lp else r.referrer_id)
        else:
            p = r.referrer_id or ""
            contact = ("•••• " + p[-4:]) if len(p) >= 4 else p
        items.append({
            "id": r.id, "code": r.code,
            "referrerKind": r.referrer_kind, "referrerContact": contact,
            "referredKind": r.referred_kind,
            "status": r.status, "reward": r.reward_amount, "payoutRef": r.payout_ref,
            "createdAt": r.created_at.isoformat() if r.created_at else None,
        })
    paid = [r for r in rows if r.status == "paid"]
    return {
        "total": len(rows), "paid": len(paid),
        "payoutTotal": sum(r.reward_amount for r in paid),
        "demoMode": settings.referral_demo_mode,
        "rewardEach": settings.referral_reward_rupees,
        "items": items,
    }


@router.get("/whatsapp/{phone}")
async def whatsapp_detail(phone: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    """Per-user detail for a WhatsApp learner: where they are, how far, quiz state."""
    s = await db.get(WhatsAppSession, phone)
    if not s:
        raise HTTPException(status_code=404, detail="WhatsApp user not found")
    from api.whatsapp_routes import _db_lessons  # lazy import avoids a circular import
    lang = s.language or "en"
    try:
        lessons = await _db_lessons(db, lang)
    except Exception:
        lessons = []
    total = len(lessons)
    idx = s.lesson_index or 0
    cur = lessons[idx] if 0 <= idx < total else None
    nudges = 0
    if isinstance(s.nudge_log, dict):
        nudges = sum((v or {}).get("n", 0) for v in s.nudge_log.values() if isinstance(v, dict))
    return {
        "type": "whatsapp",
        "name": s.name or "—", "phone": ("•••• " + phone[-4:]) if len(phone) >= 4 else phone,
        "language": lang, "stage": s.stage,
        "currentStatus": s.current_status, "goal": s.goal,
        "lesson": {"index": idx, "completed": idx, "total": total,
                   "percent": round(idx / total * 100) if total else 0,
                   "label": (cur or {}).get("label"), "title": (cur or {}).get("title")},
        "quiz": {"index": s.quiz_index or 0, "correct": s.quiz_correct or 0},
        "nudgesSent": nudges,
        "createdAt": s.created_at.isoformat() if s.created_at else None,
        "lastActive": s.last_active_at.isoformat() if s.last_active_at else None,
    }


@router.get("/whatsapp/{phone}/messages")
async def whatsapp_messages(phone: str, limit: int = 500,
                            _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    """Full WhatsApp transcript for one phone, oldest→newest. Returns the most
    recent `limit` messages (capped), then chronologically ordered for display."""
    limit = max(1, min(limit, 2000))
    total = (await db.execute(
        select(func.count()).select_from(WhatsAppMessage).where(WhatsAppMessage.phone == phone))).scalar() or 0
    # Grab the newest `limit`, then reverse to chronological order for the chat view.
    rows = (await db.execute(
        select(WhatsAppMessage).where(WhatsAppMessage.phone == phone)
        .order_by(WhatsAppMessage.created_at.desc()).limit(limit))).scalars().all()
    rows = list(reversed(rows))
    return {
        "phone": ("•••• " + phone[-4:]) if len(phone) >= 4 else phone,
        "total": total, "shown": len(rows),
        "messages": [{
            "role": m.role, "type": m.msg_type, "content": m.content,
            "at": m.created_at.isoformat() if m.created_at else None,
        } for m in rows],
    }


@router.get("/learner/{learner_id}")
async def learner_detail(learner_id: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    """Per-user detail for a web learner: completion %, exams, assignments."""
    lp = await db.get(LearnerProfile, learner_id)
    if not lp:
        raise HTTPException(status_code=404, detail="Learner not found")
    vp = (await db.execute(select(VideoProgress).where(VideoProgress.learner_id == learner_id))).scalars().all()
    last_watched = max((v.last_watched_at for v in vp if v.last_watched_at), default=None)
    # Lesson-wise progress on the SAME per-language playable lesson list WhatsApp uses
    # (only lessons that have a video), so both channels report consistently — not raw
    # Video-row counts that include empty, not-yet-uploaded lesson slots.
    from api.whatsapp_routes import _db_lessons
    try:
        lessons = await _db_lessons(db, lp.preferred_language or "en")
    except Exception:
        lessons = []
    completed_ids = {v.video_id for v in vp if v.completed}
    completed = sum(1 for les in lessons if les.get("video_id") in completed_ids)
    total_lessons = len(lessons)
    current = next((les for les in lessons if les.get("video_id") not in completed_ids), None)

    exams = (await db.execute(select(ExamAttempt).where(ExamAttempt.learner_id == learner_id)
             .order_by(ExamAttempt.attempted_at.desc()))).scalars().all()
    mod_ids = {e.module_id for e in exams if e.module_id}
    if lp.current_module_id:
        mod_ids.add(lp.current_module_id)
    mod_titles = {}
    if mod_ids:
        rows = (await db.execute(select(CourseModule.id, CourseModule.title)
                .where(CourseModule.id.in_(mod_ids)))).all()
        mod_titles = {mid: t for mid, t in rows}

    subs = (await db.execute(select(LessonAssignmentSubmission)
            .where(LessonAssignmentSubmission.learner_id == learner_id)
            .order_by(LessonAssignmentSubmission.submitted_at.desc()))).scalars().all()

    return {
        "type": "web",
        "name": lp.name, "email": lp.email, "language": lp.preferred_language, "isTest": bool(lp.is_test),
        "enrolledAt": (lp.enrollment_date or lp.created_at).isoformat() if (lp.enrollment_date or lp.created_at) else None,
        "currentModule": mod_titles.get(lp.current_module_id),
        "totalScore": lp.total_score, "certificate": bool(lp.certificate_issued),
        "lesson": {"completed": completed, "total": total_lessons,
                   "percent": round(completed / total_lessons * 100) if total_lessons else 0,
                   "label": (current or {}).get("label"), "title": (current or {}).get("title"),
                   "lastWatched": last_watched.isoformat() if last_watched else None},
        "exams": {"attempts": len(exams), "passed": sum(1 for e in exams if e.passed),
                  "bestScore": max((e.score for e in exams if e.score is not None), default=None),
                  "recent": [{"module": mod_titles.get(e.module_id, "—"), "score": e.score,
                              "passed": bool(e.passed),
                              "at": e.attempted_at.isoformat() if e.attempted_at else None} for e in exams[:6]]},
        "assignments": {"submitted": len(subs),
                        "recent": [{"lesson": a.lesson_title, "score": a.score,
                                    "at": a.submitted_at.isoformat() if a.submitted_at else None} for a in subs[:6]]},
    }


@router.get("/courses")
async def list_courses(_: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(Course).options(_FULL_TREE).order_by(Course.created_at))
    return [_course_tree(c) for c in res.scalars().all()]


class CourseBody(BaseModel):
    title: str
    description: Optional[str] = None
    thumbnail_cloudinary_id: Optional[str] = None


@router.post("/courses")
async def create_course(body: CourseBody, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    course = Course(title=body.title, description=body.description,
                    thumbnail_cloudinary_id=body.thumbnail_cloudinary_id)
    db.add(course)
    await db.commit()
    return await _tree(course.id, db)


@router.put("/courses/{course_id}")
async def update_course(course_id: str, body: CourseBody, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    course = await _load_course(course_id, db)
    course.title = body.title
    course.description = body.description
    if body.thumbnail_cloudinary_id is not None:
        course.thumbnail_cloudinary_id = body.thumbnail_cloudinary_id
    await db.commit()
    return await _tree(course_id, db)


@router.delete("/courses/{course_id}")
async def delete_course(course_id: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    module_ids = select(CourseModule.id).where(CourseModule.course_id == course_id)
    section_ids = select(Section.id).where(Section.module_id.in_(module_ids))
    video_ids = select(Video.id).where(Video.section_id.in_(section_ids))
    await db.execute(delete(VideoLanguageVariant).where(VideoLanguageVariant.video_id.in_(video_ids)))
    await db.execute(delete(VideoProgress).where(VideoProgress.video_id.in_(video_ids)))
    await db.execute(delete(Video).where(Video.section_id.in_(section_ids)))
    await db.execute(delete(Section).where(Section.module_id.in_(module_ids)))
    await db.execute(delete(CourseModule).where(CourseModule.course_id == course_id))
    await db.execute(delete(Course).where(Course.id == course_id))
    await db.commit()
    return {"deleted": True, "courseId": course_id}


# ── Modules ─────────────────────────────────────────────────────────────────────
class ModuleBody(BaseModel):
    course_id: Optional[str] = None
    title: str
    outcome: Optional[str] = None
    level: int = 1


@router.post("/modules")
async def create_module(body: ModuleBody, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    if not body.course_id:
        raise HTTPException(status_code=400, detail="course_id is required")
    order = await _next_order(db, CourseModule, CourseModule.course_id, body.course_id)
    m = CourseModule(course_id=body.course_id, title=body.title, outcome=body.outcome,
                     level=body.level, order_index=order)
    db.add(m)
    await db.commit()
    return await _tree(body.course_id, db)


@router.put("/modules/{module_id}")
async def update_module(module_id: str, body: ModuleBody, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    m = await db.get(CourseModule, module_id)
    if not m:
        raise HTTPException(status_code=404, detail="Module not found")
    m.title = body.title
    m.outcome = body.outcome
    m.level = body.level
    await db.commit()
    return await _tree(m.course_id, db)


@router.delete("/modules/{module_id}")
async def delete_module(module_id: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    m = await db.get(CourseModule, module_id)
    if not m:
        raise HTTPException(status_code=404, detail="Module not found")
    course_id = m.course_id
    section_ids = select(Section.id).where(Section.module_id == module_id)
    video_ids = select(Video.id).where(Video.section_id.in_(section_ids))
    await db.execute(delete(VideoLanguageVariant).where(VideoLanguageVariant.video_id.in_(video_ids)))
    await db.execute(delete(VideoProgress).where(VideoProgress.video_id.in_(video_ids)))
    await db.execute(delete(Video).where(Video.section_id.in_(section_ids)))
    await db.execute(delete(Section).where(Section.module_id == module_id))
    await db.execute(delete(CourseModule).where(CourseModule.id == module_id))
    await db.commit()
    return await _tree(course_id, db)


# ── Module content doc (the Teacher agent's knowledge source per module) ────────
class ContentDocBody(BaseModel):
    text: str


@router.get("/modules/{module_id}/content-doc")
async def get_module_content_doc(module_id: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    m = await db.get(CourseModule, module_id)
    if not m:
        raise HTTPException(status_code=404, detail="Module not found")
    return {"moduleId": module_id, "contentDoc": m.content_doc or ""}


@router.post("/modules/{module_id}/content-doc")
async def upload_module_content_doc(module_id: str, file: UploadFile | None = File(None), text: str | None = Form(None),
                                    _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    """Upload (or paste) the detailed sub-lesson content for a module — this
    becomes the Teacher agent's knowledge source for that module. Replaces
    whatever was there before (one doc per module)."""
    m = await db.get(CourseModule, module_id)
    if not m:
        raise HTTPException(status_code=404, detail="Module not found")
    if file is not None:
        raw = await file.read()
        if not raw:
            raise HTTPException(status_code=400, detail="The uploaded file is empty.")
        content = _extract_text(file.filename or "", raw)
    else:
        content = (text or "").strip()
    if not content:
        raise HTTPException(status_code=400, detail="No content — upload a .docx/.txt or paste the text.")
    m.content_doc = content[:200000]   # generous cap; guards against a runaway paste
    await db.commit()
    return {"moduleId": module_id, "length": len(m.content_doc)}


@router.delete("/modules/{module_id}/content-doc")
async def delete_module_content_doc(module_id: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    m = await db.get(CourseModule, module_id)
    if not m:
        raise HTTPException(status_code=404, detail="Module not found")
    m.content_doc = None
    await db.commit()
    return {"deleted": True, "moduleId": module_id}


# ── Sections ────────────────────────────────────────────────────────────────────
class SectionBody(BaseModel):
    module_id: Optional[str] = None
    title: str


@router.post("/sections")
async def create_section(body: SectionBody, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    if not body.module_id:
        raise HTTPException(status_code=400, detail="module_id is required")
    m = await db.get(CourseModule, body.module_id)
    if not m:
        raise HTTPException(status_code=404, detail="Module not found")
    order = await _next_order(db, Section, Section.module_id, body.module_id)
    s = Section(module_id=body.module_id, title=body.title, order_index=order)
    db.add(s)
    await db.commit()
    return await _tree(m.course_id, db)


@router.put("/sections/{section_id}")
async def update_section(section_id: str, body: SectionBody, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    s = await db.get(Section, section_id)
    if not s:
        raise HTTPException(status_code=404, detail="Section not found")
    s.title = body.title
    m = await db.get(CourseModule, s.module_id)
    await db.commit()
    return await _tree(m.course_id, db)


@router.delete("/sections/{section_id}")
async def delete_section(section_id: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    s = await db.get(Section, section_id)
    if not s:
        raise HTTPException(status_code=404, detail="Section not found")
    m = await db.get(CourseModule, s.module_id)
    video_ids = select(Video.id).where(Video.section_id == section_id)
    await db.execute(delete(VideoLanguageVariant).where(VideoLanguageVariant.video_id.in_(video_ids)))
    await db.execute(delete(VideoProgress).where(VideoProgress.video_id.in_(video_ids)))
    await db.execute(delete(Video).where(Video.section_id == section_id))
    await db.execute(delete(Section).where(Section.id == section_id))
    await db.commit()
    return await _tree(m.course_id, db)


# ── Videos (lessons) ────────────────────────────────────────────────────────────
class VideoBody(BaseModel):
    section_id: Optional[str] = None
    title: str
    cloudinary_public_id: Optional[str] = None
    duration_seconds: Optional[int] = None


@router.post("/videos")
async def create_video(body: VideoBody, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    if not body.section_id:
        raise HTTPException(status_code=400, detail="section_id is required")
    s = await db.get(Section, body.section_id)
    if not s:
        raise HTTPException(status_code=404, detail="Section not found")
    order = await _next_order(db, Video, Video.section_id, body.section_id)
    v = Video(section_id=body.section_id, title=body.title,
              cloudinary_public_id=body.cloudinary_public_id,
              duration_seconds=body.duration_seconds, order_index=order)
    db.add(v)
    await db.commit()
    m = await db.get(CourseModule, s.module_id)
    return await _tree(m.course_id, db)


async def _warm_derivative(public_id: str) -> None:
    """Request the delivery transform once so Cloudinary generates + caches it
    (Cloudinary-side transcode — no local processing). Makes the first WhatsApp
    send fast instead of hitting a cold transcode. Same transform string as the
    WhatsApp bot + web, so all channels reuse this one cached derivative. Warms
    ONE video (the one just uploaded) — not a blanket sweep."""
    from api.whatsapp_routes import VIDEO_TRANSFORM  # single source of the string
    if not settings.cloudinary_cloud_name or not public_id:
        return
    url = (f"https://res.cloudinary.com/{settings.cloudinary_cloud_name}"
           f"/video/upload/{VIDEO_TRANSFORM}/{public_id}.mp4")
    try:
        async with httpx.AsyncClient(timeout=120, follow_redirects=True) as h:
            r = await h.get(url)
        print(f"warm {public_id}: {r.status_code}")
    except Exception as e:
        print(f"⚠ warm error {public_id}: {e}")


@router.put("/videos/{video_id}")
async def update_video(video_id: str, body: VideoBody, background_tasks: BackgroundTasks, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    v = await db.get(Video, video_id)
    if not v:
        raise HTTPException(status_code=404, detail="Video not found")
    v.title = body.title
    new_upload = body.cloudinary_public_id is not None and body.cloudinary_public_id != v.cloudinary_public_id
    if body.cloudinary_public_id is not None:
        v.cloudinary_public_id = body.cloudinary_public_id
    if body.duration_seconds is not None:
        v.duration_seconds = body.duration_seconds
    s = await db.get(Section, v.section_id)
    m = await db.get(CourseModule, s.module_id)
    if new_upload and v.cloudinary_public_id:
        background_tasks.add_task(_warm_derivative, v.cloudinary_public_id)
    await db.commit()
    return await _tree(m.course_id, db)


@router.delete("/videos/{video_id}")
async def delete_video(video_id: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    v = await db.get(Video, video_id)
    if not v:
        raise HTTPException(status_code=404, detail="Video not found")
    s = await db.get(Section, v.section_id)
    m = await db.get(CourseModule, s.module_id)
    await db.execute(delete(VideoLanguageVariant).where(VideoLanguageVariant.video_id == video_id))
    await db.execute(delete(VideoProgress).where(VideoProgress.video_id == video_id))
    await db.execute(delete(Video).where(Video.id == video_id))
    await db.commit()
    return await _tree(m.course_id, db)


# ── Per-language video variants ──────────────────────────────────────────────────
class VariantBody(BaseModel):
    language: str
    cloudinary_public_id: str
    duration_seconds: Optional[int] = None


@router.put("/videos/{video_id}/variant")
async def upsert_variant(video_id: str, body: VariantBody, background_tasks: BackgroundTasks, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    if body.language not in SUPPORTED_LANGUAGES:
        raise HTTPException(status_code=400, detail=f"Unsupported language. Supported: {sorted(SUPPORTED_LANGUAGES)}")
    v = await db.get(Video, video_id)
    if not v:
        raise HTTPException(status_code=404, detail="Video not found")
    res = await db.execute(select(VideoLanguageVariant).where(
        VideoLanguageVariant.video_id == video_id, VideoLanguageVariant.language == body.language))
    variant = res.scalar_one_or_none()
    if variant:
        new_upload = body.cloudinary_public_id != variant.cloudinary_public_id
        variant.cloudinary_public_id = body.cloudinary_public_id
        if body.duration_seconds is not None:
            variant.duration_seconds = body.duration_seconds
    else:
        new_upload = True
        variant = VideoLanguageVariant(video_id=video_id, language=body.language,
                                       cloudinary_public_id=body.cloudinary_public_id,
                                       duration_seconds=body.duration_seconds)
        db.add(variant)
    if new_upload and variant.cloudinary_public_id:
        background_tasks.add_task(_warm_derivative, variant.cloudinary_public_id)
    await db.commit()
    return {"ok": True, "videoId": video_id, "language": body.language}


@router.delete("/videos/{video_id}/variant/{language}")
async def delete_variant(video_id: str, language: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(VideoLanguageVariant).where(
        VideoLanguageVariant.video_id == video_id, VideoLanguageVariant.language == language))
    variant = res.scalar_one_or_none()
    if not variant:
        raise HTTPException(status_code=404, detail="Variant not found")
    await db.delete(variant)
    await db.commit()
    return {"deleted": True, "videoId": video_id, "language": language}


# ── Cloudinary signed direct upload ───────────────────────────────────────────────
class SignatureBody(BaseModel):
    folder: Optional[str] = None
    resource_type: Optional[str] = "video"   # 'video' | 'image' | 'auto'


@router.post("/cloudinary/signature")
async def cloudinary_signature(body: SignatureBody, _: bool = Depends(require_admin)):
    """Return params for a signed direct browser→Cloudinary upload. The api_secret
    never leaves the server — only the derived signature does."""
    if not settings.cloudinary_api_key or not settings.cloudinary_api_secret:
        raise HTTPException(status_code=400,
                            detail="Cloudinary upload is not configured on the server (missing API key/secret).")
    ts = int(time.time())
    folder = body.folder or "cosmoplex/lessons"
    rtype = (body.resource_type or "video").lower()
    if rtype not in ("video", "image", "auto"):
        raise HTTPException(status_code=400, detail="resource_type must be video, image, or auto")
    # resource_type is part of the URL, not a signed param — only folder+timestamp
    # are signed. Sign the params Cloudinary will receive (alphabetical), then append secret.
    to_sign = f"folder={folder}&timestamp={ts}{settings.cloudinary_api_secret}"
    signature = hashlib.sha1(to_sign.encode()).hexdigest()
    return {
        "timestamp": ts,
        "signature": signature,
        "apiKey": settings.cloudinary_api_key,
        "cloudName": settings.cloudinary_cloud_name,
        "folder": folder,
        "uploadUrl": f"https://api.cloudinary.com/v1_1/{settings.cloudinary_cloud_name}/{rtype}/upload",
    }


# ── Quiz bank (per lesson) ───────────────────────────────────────────────────
class QuizBody(BaseModel):
    question: dict                 # {"en": "...", "hi": "...", ...}
    options: dict                  # {"en": ["a","b","c","d"], ...}
    correct_index: int


def _quiz_dict(q: QuizQuestion) -> dict:
    return {"id": q.id, "question": q.question, "options": q.options,
            "correctIndex": q.correct_index, "orderIndex": q.order_index}


@router.get("/videos/{video_id}/quizzes")
async def list_quizzes(video_id: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(QuizQuestion).where(QuizQuestion.video_id == video_id).order_by(QuizQuestion.order_index))
    return [_quiz_dict(q) for q in res.scalars().all()]


@router.post("/videos/{video_id}/quizzes")
async def create_quiz(video_id: str, body: QuizBody, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    if not await db.get(Video, video_id):
        raise HTTPException(status_code=404, detail="Video not found")
    if not body.question.get("en"):
        raise HTTPException(status_code=400, detail="English question is required")
    order = await _next_order(db, QuizQuestion, QuizQuestion.video_id, video_id)
    q = QuizQuestion(video_id=video_id, question=body.question, options=body.options,
                     correct_index=body.correct_index, order_index=order)
    db.add(q)
    await db.commit()
    return _quiz_dict(q)


@router.put("/quizzes/{quiz_id}")
async def update_quiz(quiz_id: str, body: QuizBody, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    q = await db.get(QuizQuestion, quiz_id)
    if not q:
        raise HTTPException(status_code=404, detail="Question not found")
    q.question = body.question
    q.options = body.options
    q.correct_index = body.correct_index
    await db.commit()
    return _quiz_dict(q)


@router.delete("/quizzes/{quiz_id}")
async def delete_quiz(quiz_id: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    q = await db.get(QuizQuestion, quiz_id)
    if not q:
        raise HTTPException(status_code=404, detail="Question not found")
    await db.delete(q)
    await db.commit()
    return {"deleted": True, "id": quiz_id}


# ── Assignment bank (per lesson) ─────────────────────────────────────────────
class AssignmentBody(BaseModel):
    question: dict                 # {"en": "...", "hi": "...", ...}
    rubric: str


def _assign_dict(a: AssignmentPrompt) -> dict:
    return {"id": a.id, "question": a.question, "rubric": a.rubric, "orderIndex": a.order_index}


@router.get("/videos/{video_id}/assignments")
async def list_assignments(video_id: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(AssignmentPrompt).where(AssignmentPrompt.video_id == video_id).order_by(AssignmentPrompt.order_index))
    return [_assign_dict(a) for a in res.scalars().all()]


@router.post("/videos/{video_id}/assignments")
async def create_assignment(video_id: str, body: AssignmentBody, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    if not await db.get(Video, video_id):
        raise HTTPException(status_code=404, detail="Video not found")
    if not body.question.get("en"):
        raise HTTPException(status_code=400, detail="English question is required")
    order = await _next_order(db, AssignmentPrompt, AssignmentPrompt.video_id, video_id)
    a = AssignmentPrompt(video_id=video_id, question=body.question, rubric=body.rubric, order_index=order)
    db.add(a)
    await db.commit()
    return _assign_dict(a)


@router.put("/assignments/{assignment_id}")
async def update_assignment(assignment_id: str, body: AssignmentBody, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    a = await db.get(AssignmentPrompt, assignment_id)
    if not a:
        raise HTTPException(status_code=404, detail="Assignment not found")
    a.question = body.question
    a.rubric = body.rubric
    await db.commit()
    return _assign_dict(a)


@router.delete("/assignments/{assignment_id}")
async def delete_assignment(assignment_id: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    a = await db.get(AssignmentPrompt, assignment_id)
    if not a:
        raise HTTPException(status_code=404, detail="Assignment not found")
    await db.delete(a)
    await db.commit()
    return {"deleted": True, "id": assignment_id}


# ── WhatsApp intro videos (onboarding) ─────────────────────────────────────────
INTRO_LANGUAGES = {"default"} | SUPPORTED_LANGUAGES


class IntroVideoBody(BaseModel):
    cloudinary_public_id: str
    duration_seconds: Optional[int] = None


def _intro_dict(iv: IntroVideo) -> dict:
    return {"language": iv.language, "cloudinaryPublicId": iv.cloudinary_public_id,
            "durationSeconds": iv.duration_seconds}


@router.get("/intro-videos")
async def list_intro_videos(_: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(IntroVideo))
    return [_intro_dict(iv) for iv in res.scalars().all()]


@router.put("/intro-videos/{language}")
async def set_intro_video(language: str, body: IntroVideoBody,
                          _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    if language not in INTRO_LANGUAGES:
        raise HTTPException(status_code=400, detail=f"language must be one of {sorted(INTRO_LANGUAGES)}")
    if not body.cloudinary_public_id.strip():
        raise HTTPException(status_code=400, detail="cloudinary_public_id is required")
    iv = await db.get(IntroVideo, language)
    if iv is None:
        iv = IntroVideo(language=language)
        db.add(iv)
    iv.cloudinary_public_id = body.cloudinary_public_id.strip()
    iv.duration_seconds = body.duration_seconds
    await db.commit()
    return _intro_dict(iv)


@router.delete("/intro-videos/{language}")
async def delete_intro_video(language: str, _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    iv = await db.get(IntroVideo, language)
    if not iv:
        raise HTTPException(status_code=404, detail="Intro video not set for this language")
    await db.delete(iv)
    await db.commit()
    return {"deleted": True, "language": language}


# ── Pre-sale marketing assets (drip media for stuck-at-signup users) ───────────
# One photo OR video per (day-bucket, language). Days = idle-days before it fires.
MARKETING_DAYS = [1, 2, 3, 7]


class MarketingAssetBody(BaseModel):
    # All optional — a PUT patches only the fields it includes (send null to clear one).
    image_public_id: Optional[str] = None
    video_public_id: Optional[str] = None
    video_duration_seconds: Optional[int] = None
    text: Optional[str] = None


def _marketing_dict(a: MarketingAsset) -> dict:
    return {"day": a.day, "language": a.language,
            "imagePublicId": a.image_public_id, "videoPublicId": a.video_public_id,
            "videoDurationSeconds": a.video_duration_seconds, "text": a.text}


@router.get("/marketing-assets")
async def list_marketing_assets(_: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(MarketingAsset))
    return {"days": MARKETING_DAYS, "languages": sorted(SUPPORTED_LANGUAGES),
            "items": [_marketing_dict(a) for a in res.scalars().all()]}


@router.put("/marketing-assets/{day}/{language}")
async def set_marketing_asset(day: int, language: str, body: MarketingAssetBody,
                              _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    if day not in MARKETING_DAYS:
        raise HTTPException(status_code=400, detail=f"day must be one of {MARKETING_DAYS}")
    if language not in SUPPORTED_LANGUAGES:
        raise HTTPException(status_code=400, detail=f"language must be one of {sorted(SUPPORTED_LANGUAGES)}")
    key = f"{day}_{language}"
    a = await db.get(MarketingAsset, key)
    if a is None:
        a = MarketingAsset(id=key, day=day, language=language)
        db.add(a)
    # Patch only the fields present in the request (exclude_unset), so each of the
    # three fields — photo, video, text — can be set or cleared independently.
    patch = body.model_dump(exclude_unset=True)
    for field in ("image_public_id", "video_public_id", "video_duration_seconds", "text"):
        if field in patch:
            val = patch[field]
            if isinstance(val, str):
                val = val.strip() or None
            setattr(a, field, val)
    await db.commit()
    return _marketing_dict(a)


@router.delete("/marketing-assets/{day}/{language}")
async def delete_marketing_asset(day: int, language: str,
                                 _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    a = await db.get(MarketingAsset, f"{day}_{language}")
    if not a:
        raise HTTPException(status_code=404, detail="No asset for that day/language")
    await db.delete(a)
    await db.commit()
    return {"deleted": True, "day": day, "language": language}


# ── Bulk import: upload a doc of questions → Claude extracts + translates ───────
# Lets an admin drop a .docx/.txt (or paste text) instead of hand-entering each
# question in 6 languages. Claude pulls out the questions and translates every
# one into all 6 languages, then they're appended to the lesson's bank.
BULK_MODEL = "claude-sonnet-4-6"   # quality matters across languages; admin-only, infrequent
BULK_LANGS = ["en", "hi", "mr", "te", "ta", "kn"]

QUIZ_SYS = """You extract multiple-choice quiz questions from a document and translate them.
The text may be messy (copied from Word, tables, numbered lists). Find EVERY multiple-choice question.
For each: the question text, its 2-4 answer options, and which option is correct.
Translate the question and EVERY option into: English(en), Hindi(hi), Marathi(mr), Telugu(te), Tamil(ta), Kannada(kn). Keep translations faithful and natural for Indian learners; keep technical AI terms recognizable.
Return ONLY strict JSON (no markdown, no commentary) shaped exactly:
{"questions":[{"question":{"en":"","hi":"","mr":"","te":"","ta":"","kn":""},"options":{"en":["",""],"hi":["",""],"mr":["",""],"te":["",""],"ta":["",""],"kn":["",""]},"correct_index":0}]}
Rules:
- Extract EVERY question present — do NOT skip, merge, renumber, summarise, or deduplicate any. If the text has 20 questions, return all 20.
- correct_index is 0-based into the options arrays.
- Every language's options array MUST have the same number of items, in the same order, as English.
- Detect the correct answer from any marker in the source (*, "Answer:", "✓", "Correct", "[B]", bold). If none, choose the best answer.
- Do not invent extra questions. Output only the JSON object."""

ASSIGN_SYS = """You extract open-ended assignment questions from a document and translate them.
Find EVERY assignment/task prompt (written answers, NOT multiple choice).
Translate each question into: English(en), Hindi(hi), Marathi(mr), Telugu(te), Tamil(ta), Kannada(kn). Faithful and natural.
For each assignment also produce a short grading "rubric" in ENGLISH (language-neutral) for an AI grader — use the document's grading criteria if given, otherwise write a concise 1-2 sentence rubric from the question.
Return ONLY strict JSON (no markdown, no commentary) shaped exactly:
{"assignments":[{"question":{"en":"","hi":"","mr":"","te":"","ta":"","kn":""},"rubric":""}]}
Output only the JSON object."""


def _extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".docx"):
        try:
            import docx
        except ImportError:
            raise HTTPException(status_code=500,
                detail="Server can't read .docx (python-docx missing). Paste the text or upload a .txt.")
        try:
            doc = docx.Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append("\t".join(c.text for c in row.cells))
            return "\n".join(parts).strip()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Couldn't read that .docx: {e}")
    if name.endswith(".doc"):
        raise HTTPException(status_code=400,
            detail="Old .doc isn't supported — save as .docx or .txt, or paste the text.")
    return data.decode("utf-8", errors="ignore").strip()  # txt / csv / md / other


async def _claude_json(system: str, user: str, max_tokens: int = 8000) -> dict:
    if not settings.anthropic_api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on the server.")
    from anthropic import AsyncAnthropic
    client = AsyncAnthropic(api_key=settings.anthropic_api_key)
    try:
        resp = await client.messages.create(
            model=BULK_MODEL, max_tokens=max_tokens,
            system=system, messages=[{"role": "user", "content": user}],
        )
        raw = (resp.content[0].text or "").strip()
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"AI extraction failed: {e}")
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        i, j = raw.find("{"), raw.rfind("}")   # tolerate stray prose / code fences
        if i != -1 and j > i:
            try:
                return json.loads(raw[i:j + 1])
            except json.JSONDecodeError:
                pass
        raise HTTPException(status_code=422,
            detail="Couldn't parse questions from that document — it may be too long or unclear. "
                   "Try fewer questions or a cleaner layout.")


def _clean_quiz(items) -> list[dict]:
    out: list[dict] = []
    for it in (items or []):
        q, opts = (it.get("question") or {}), (it.get("options") or {})
        if not isinstance(q, dict) or not isinstance(opts, dict):
            continue
        en_q = (q.get("en") or "").strip()
        en_opts = opts.get("en")
        if not en_q or not isinstance(en_opts, list) or len(en_opts) < 2:
            continue
        n = len(en_opts)
        try:
            ci = max(0, min(int(it.get("correct_index", 0)), n - 1))
        except (TypeError, ValueError):
            ci = 0
        clean_q = {"en": en_q}
        clean_opts = {"en": [str(o).strip() for o in en_opts]}
        for lang in BULK_LANGS[1:]:
            lv_q, lv_o = (q.get(lang) or "").strip(), opts.get(lang)
            if lv_q and isinstance(lv_o, list) and len(lv_o) == n and all(str(x).strip() for x in lv_o):
                clean_q[lang] = lv_q
                clean_opts[lang] = [str(o).strip() for o in lv_o]
        out.append({"question": clean_q, "options": clean_opts, "correct_index": ci})
    return out


def _clean_assignments(items) -> list[dict]:
    out: list[dict] = []
    for it in (items or []):
        q = it.get("question") or {}
        if not isinstance(q, dict):
            continue
        en_q = (q.get("en") or "").strip()
        if not en_q:
            continue
        rubric = (it.get("rubric") or "").strip() or (
            "Evaluate whether the answer correctly and clearly addresses the question and "
            "shows understanding of the concept.")
        clean_q = {"en": en_q}
        for lang in BULK_LANGS[1:]:
            lv = (q.get(lang) or "").strip()
            if lv:
                clean_q[lang] = lv
        out.append({"question": clean_q, "rubric": rubric})
    return out


async def _bulk_content(file: UploadFile | None, text: str | None) -> str:
    if file is not None:
        data = await file.read()
        if not data:
            raise HTTPException(status_code=400, detail="The uploaded file is empty.")
        content = _extract_text(file.filename or "", data)
    else:
        content = (text or "").strip()
    if not content:
        raise HTTPException(status_code=400, detail="No content — upload a .docx/.txt or paste the questions.")
    return content[:60000]   # guard against a huge doc blowing the token budget


def _chunk_questions(text: str, per_chunk: int = 5) -> list[str]:
    """Split a questions doc into batches of ~per_chunk numbered items. Each batch
    is extracted + translated into 6 languages in its own Claude call, keeping the
    output within the token limit (translating 20 questions at once overflows and
    truncates the JSON). Falls back to one chunk if items can't be detected."""
    lines = text.split("\n")
    starts = [i for i, l in enumerate(lines) if re.match(r"\s*(?:Q\s*)?\d+\s*[.)\:]", l)]
    if len(starts) < 2:
        return [text]
    header = "\n".join(lines[:starts[0]]).strip()
    blocks = ["\n".join(lines[starts[k]: (starts[k + 1] if k + 1 < len(starts) else len(lines))])
              for k in range(len(starts))]
    chunks = []
    for k in range(0, len(blocks), per_chunk):
        body = "\n".join(blocks[k:k + per_chunk])
        chunks.append((header + "\n\n" + body) if header else body)
    return chunks


async def _extract_items(system: str, content: str, key: str) -> list[dict]:
    """Extract + translate across batches (run in parallel) and merge the raw item
    lists under `key` ('questions' | 'assignments')."""
    chunks = _chunk_questions(content)
    results = await asyncio.gather(*[_claude_json(system, ch) for ch in chunks],
                                   return_exceptions=True)
    merged: list[dict] = []
    for r in results:
        if isinstance(r, dict):
            merged.extend(r.get(key) or [])
    if not merged:
        raise HTTPException(status_code=422,
            detail="Couldn't parse questions from that document — try a cleaner layout, "
                   "or split it into a couple of smaller uploads.")
    return merged


# ── Deterministic MCQ parsing (no AI guessing) + translation-only step ──────────
_Q_RE = re.compile(r"^(?:Q\s*)?\d+\s*[.)\:]\s*(.+)$")
_OPT_RE = re.compile(r"^\(?([a-eA-E])[.)]\s*(.+)$")
_CORRECT_RE = re.compile(r"✓|✔|\bcorrect\b", re.I)

TRANSLATE_SYS = """You translate multiple-choice quiz questions from English into 5 Indian languages.
Input: a JSON array of questions, each {"q":"<english>","options":["<en>", ...]}.
Translate each question and EACH option into Hindi(hi), Marathi(mr), Telugu(te), Tamil(ta), Kannada(kn).
Faithful, natural for Indian learners; keep technical AI terms recognizable. Do NOT change the English.
Return ONLY strict JSON: {"translations":[{"q":{"hi":"","mr":"","te":"","ta":"","kn":""},"options":{"hi":["",...],"mr":[...],"te":[...],"ta":[...],"kn":[...]}}]}
The array MUST be the same length and order as the input; each options array MUST have the same number of items, in the same order. Do not add, drop, or reorder anything."""


def _parse_mcq(text: str) -> list[dict]:
    """Parse numbered MCQs (question + a/b/c/d options; correct marked with ✓ or
    'Correct') deterministically — reliable, no AI interpretation. Returns English
    items {q, options, correct_index}. Empty if the doc isn't in this structure."""
    out: list[dict] = []
    cur: dict | None = None
    for raw in text.split("\n"):
        s = raw.strip()
        if not s:
            continue
        opt = _OPT_RE.match(s)
        q = _Q_RE.match(s) if not opt else None
        if q:
            if cur and len(cur["options"]) >= 2:
                out.append(cur)
            qt = re.sub(r"\s*\[[A-Za-z]\]\s*$", "", q.group(1)).strip()  # drop [B]/[I]/[A] difficulty tags
            cur = {"q": qt, "options": [], "correct_index": 0}
        elif opt and cur is not None:
            body = opt.group(2)
            is_correct = bool(_CORRECT_RE.search(body))
            body = re.sub(r"\s*[✓✔].*$", "", body)                 # strip "✓ Correct"
            body = re.sub(r"\s*[-—]?\s*\bcorrect\b\s*$", "", body, flags=re.I).strip()
            cur["options"].append(body)
            if is_correct:
                cur["correct_index"] = len(cur["options"]) - 1
    if cur and len(cur["options"]) >= 2:
        out.append(cur)
    return out


async def _translate_quiz(english: list[dict]) -> list[dict]:
    """Translate parsed English MCQs into all languages (AI does ONLY translation).
    English + correct answer are authoritative; a failed batch just leaves those
    items English-only."""
    batches = [english[i:i + 6] for i in range(0, len(english), 6)]
    results = await asyncio.gather(
        *[_claude_json(TRANSLATE_SYS,
                       json.dumps([{"q": it["q"], "options": it["options"]} for it in b], ensure_ascii=False))
          for b in batches],
        return_exceptions=True,
    )
    final: list[dict] = []
    for b, r in zip(batches, results):
        trs = (r.get("translations") if isinstance(r, dict) else None) or []
        for i, en in enumerate(b):
            n = len(en["options"])
            question = {"en": en["q"]}
            options = {"en": [str(o).strip() for o in en["options"]]}
            tr = trs[i] if i < len(trs) else {}
            tq, topts = (tr.get("q") or {}), (tr.get("options") or {})
            for lang in ("hi", "mr", "te", "ta", "kn"):
                lq, lo = (tq.get(lang) or "").strip(), topts.get(lang)
                if lq and isinstance(lo, list) and len(lo) == n and all(str(x).strip() for x in lo):
                    question[lang] = lq
                    options[lang] = [str(x).strip() for x in lo]
            final.append({"question": question, "options": options, "correct_index": en["correct_index"]})
    return final


TRANSLATE_ASSIGN_SYS = """You translate open-ended assignment prompts from English into 5 Indian languages.
Input: a JSON array of {"q":"<english prompt>"}.
Translate each prompt into Hindi(hi), Marathi(mr), Telugu(te), Tamil(ta), Kannada(kn). Faithful, natural for Indian learners; keep technical AI terms recognizable. Do NOT change the English.
Return ONLY strict JSON: {"translations":[{"q":{"hi":"","mr":"","te":"","ta":"","kn":""}}]}
The array MUST be the same length and order as the input. Do not add, drop, or reorder anything."""

DEFAULT_RUBRIC = ("Evaluate whether the answer correctly and clearly addresses the prompt and shows "
                  "genuine understanding of the concept, explained in the learner's own words.")


def _parse_assignments(text: str) -> list[dict]:
    """Parse numbered assignment prompts (1., 2., Q1) …) deterministically. A prompt
    may span multiple lines (joined). Returns English items {q}. Empty if not numbered."""
    lines = text.split("\n")
    starts = [i for i, l in enumerate(lines) if re.match(r"^\s*(?:Q\s*)?\d+\s*[.)\:]\s*\S", l)]
    if not starts:
        return []
    out: list[dict] = []
    for k in range(len(starts)):
        s = starts[k]
        e = starts[k + 1] if k + 1 < len(starts) else len(lines)
        block = " ".join(x.strip() for x in lines[s:e] if x.strip())
        block = re.sub(r"^\s*(?:Q\s*)?\d+\s*[.)\:]\s*", "", block).strip()
        block = re.sub(r"\s*\[[^\]]{1,20}\]\s*$", "", block).strip()   # drop trailing [tags]
        if len(block) >= 8:
            out.append({"q": block})
    return out


_TASK_TITLE_RE = re.compile(r"^[A-Za-z]?\d+\s*[—–\-]\s*(.+)$")


def _iter_docx_blocks(doc):
    """Yield ('p', Paragraph) / ('tbl', Table) in TRUE document order (python-docx's
    doc.paragraphs / doc.tables are separate flat lists that lose interleaving)."""
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield ("p", Paragraph(child, doc))
        elif tag == "tbl":
            yield ("tbl", Table(child, doc))


def _table_is_rubric(t) -> bool:
    if not t.rows:
        return False
    header = [c.text.strip().lower() for c in t.rows[0].cells]
    return any("submit" in h for h in header) and any("pass" in h for h in header)


def _parse_task_assigner_pack(data: bytes) -> list[dict]:
    """Deterministic parser for the 'Task Assigner content pack' template: a title
    line like 'A1 — Spot It In Your Field', a short label ('Task Assigner message:'),
    the prompt paragraph, then Covers/Format + Submit/Pass/Fail-nudge tables (in any
    order, possibly interleaved). Builds a real grading rubric from the Pass/Fail-nudge
    cells instead of a generic placeholder. Returns [] if the doc isn't this template."""
    try:
        import docx
    except ImportError:
        return []
    try:
        doc = docx.Document(io.BytesIO(data))
        blocks = list(_iter_docx_blocks(doc))
    except Exception:
        return []

    title_idx = [i for i, (k, o) in enumerate(blocks) if k == "p" and _TASK_TITLE_RE.match(o.text.strip())]
    if len(title_idx) < 2:      # need at least 2 to be confident this is the template, not a stray line
        return []

    out: list[dict] = []
    for ti, start in enumerate(title_idx):
        end = title_idx[ti + 1] if ti + 1 < len(title_idx) else len(blocks)
        span = blocks[start + 1:end]

        # The prompt is the longest paragraph in the span — labels/headers are short.
        para_texts = [o.text.strip() for k, o in span if k == "p" and o.text.strip()]
        prompt = max(para_texts, key=len) if para_texts else None
        if not prompt or len(prompt) < 8:
            continue

        pass_txt = fail_txt = None
        for k, o in span:
            if k == "tbl" and _table_is_rubric(o):
                header = [c.text.strip().lower() for c in o.rows[0].cells]
                data_row = o.rows[1].cells if len(o.rows) > 1 else None
                if data_row:
                    for hi, h in enumerate(header):
                        if "pass" in h:
                            pass_txt = data_row[hi].text.strip()
                        elif "fail" in h:
                            fail_txt = data_row[hi].text.strip()
                break

        rubric = None
        if pass_txt and fail_txt:
            rubric = f"Pass if: {pass_txt}. If not met, note: {fail_txt}"
        elif pass_txt:
            rubric = f"Pass if: {pass_txt}"
        out.append({"q": prompt, "rubric": rubric})
    return out


async def _translate_assignments(english: list[dict]) -> list[dict]:
    """Translate parsed English prompts into all languages (AI does ONLY translation).
    A failed batch leaves those items English-only. Uses each item's own `rubric` if
    given (e.g. built from a doc's Pass/Fail table); otherwise a sensible default."""
    batches = [english[i:i + 8] for i in range(0, len(english), 8)]
    results = await asyncio.gather(
        *[_claude_json(TRANSLATE_ASSIGN_SYS, json.dumps([{"q": it["q"]} for it in b], ensure_ascii=False))
          for b in batches],
        return_exceptions=True,
    )
    final: list[dict] = []
    for b, r in zip(batches, results):
        trs = (r.get("translations") if isinstance(r, dict) else None) or []
        for i, en in enumerate(b):
            question = {"en": en["q"]}
            tr = trs[i] if i < len(trs) else {}
            tq = tr.get("q") or {}
            for lang in ("hi", "mr", "te", "ta", "kn"):
                lq = (tq.get(lang) or "").strip()
                if lq:
                    question[lang] = lq
            final.append({"question": question, "rubric": (en.get("rubric") or "").strip() or DEFAULT_RUBRIC})
    return final


# ── Gap-fill: re-translate any item a batch left English-only ──────────────────
# The batch translate is best-effort — a failed/truncated batch leaves some items
# with missing languages, which used to be stored English-only silently. This pass
# re-translates just those items, one at a time (small output, no truncation), so
# nothing is ever silently English-only. A failure here leaves the item as-is
# (English still works) rather than aborting the whole import.
_GAP_QUIZ_SYS = ("Translate one multiple-choice question from English into Hindi(hi), Marathi(mr), "
                 "Telugu(te), Tamil(ta), Kannada(kn). Faithful, natural for Indian learners; keep "
                 "technical AI terms recognizable. Keep each options array the SAME length and order. "
                 'Return ONLY strict JSON: {"hi":{"q":"","options":["",...]},"mr":{...},"te":{...},"ta":{...},"kn":{...}}')
_GAP_ASSIGN_SYS = ("Translate one open-ended assignment prompt from English into Hindi(hi), Marathi(mr), "
                   "Telugu(te), Tamil(ta), Kannada(kn). Faithful, natural; keep technical AI terms. "
                   'Return ONLY strict JSON: {"hi":"","mr":"","te":"","ta":"","kn":""}')


async def _fill_quiz_gaps(items: list[dict]) -> list[dict]:
    async def fill(it):
        q, o = it["question"], it["options"]
        n = len(o.get("en") or [])
        if n < 2 or all(q.get(l) for l in BULK_LANGS[1:]):
            return it
        try:
            tr = await _claude_json(_GAP_QUIZ_SYS, json.dumps({"q": q["en"], "options": o["en"]}, ensure_ascii=False))
        except HTTPException:
            return it
        for lang in BULK_LANGS[1:]:
            if q.get(lang):
                continue
            d = tr.get(lang) or {}
            lq, lo = (d.get("q") or "").strip(), d.get("options")
            if lq and isinstance(lo, list) and len(lo) == n and all(str(x).strip() for x in lo):
                q[lang] = lq
                o[lang] = [str(x).strip() for x in lo]
        return it
    return list(await asyncio.gather(*[fill(it) for it in items]))


async def _fill_assign_gaps(items: list[dict]) -> list[dict]:
    async def fill(it):
        q = it["question"]
        if all(q.get(l) for l in BULK_LANGS[1:]):
            return it
        try:
            tr = await _claude_json(_GAP_ASSIGN_SYS, q["en"])
        except HTTPException:
            return it
        for lang in BULK_LANGS[1:]:
            if not q.get(lang) and (tr.get(lang) or "").strip():
                q[lang] = tr[lang].strip()
        return it
    return list(await asyncio.gather(*[fill(it) for it in items]))


@router.post("/videos/{video_id}/quizzes/bulk")
async def bulk_quizzes(video_id: str, file: UploadFile | None = File(None), text: str | None = Form(None),
                       replace: bool = Form(False),
                       _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    """Upload a doc / paste text of MCQs → extract + translate. Appends by default;
    `replace=true` clears this lesson's existing quiz bank first (clean re-import)."""
    if not await db.get(Video, video_id):
        raise HTTPException(status_code=404, detail="Video not found")
    content = await _bulk_content(file, text)
    parsed = _parse_mcq(content)
    if parsed:
        # Deterministic extraction (exact count/options/answer) + AI translation only.
        items = _clean_quiz(await _translate_quiz(parsed))
    else:
        # Unstructured doc → let the AI find the questions too.
        items = _clean_quiz(await _extract_items(QUIZ_SYS, content, "questions"))
    items = await _fill_quiz_gaps(items)   # retry any language a batch missed
    if not items:
        raise HTTPException(status_code=422, detail="No multiple-choice questions found in that document.")
    if replace:
        await db.execute(delete(QuizQuestion).where(QuizQuestion.video_id == video_id))
    order = 0 if replace else await _next_order(db, QuizQuestion, QuizQuestion.video_id, video_id)
    created = [QuizQuestion(video_id=video_id, question=it["question"], options=it["options"],
                            correct_index=it["correct_index"], order_index=order + i)
               for i, it in enumerate(items)]
    for q in created:
        db.add(q)
    await db.commit()
    return {"added": len(created), "replaced": replace, "items": [_quiz_dict(q) for q in created]}


@router.post("/videos/{video_id}/assignments/bulk")
async def bulk_assignments(video_id: str, file: UploadFile | None = File(None), text: str | None = Form(None),
                           replace: bool = Form(False),
                           _: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    """Upload a doc / paste text of assignment prompts → extract + translate. Appends
    by default; `replace=true` clears this lesson's existing assignment bank first."""
    if not await db.get(Video, video_id):
        raise HTTPException(status_code=404, detail="Video not found")

    parsed: list[dict] = []
    if file is not None:
        raw = await file.read()
        if not raw:
            raise HTTPException(status_code=400, detail="The uploaded file is empty.")
        if (file.filename or "").lower().endswith(".docx"):
            # Try the structured "Task Assigner pack" template first (title + label +
            # prompt + Covers/Format + Submit/Pass/Fail-nudge tables) — builds a real
            # rubric from the doc instead of a generic placeholder.
            parsed = _parse_task_assigner_pack(raw)
        content = _extract_text(file.filename or "", raw)
    else:
        content = (text or "").strip()
    if not content and not parsed:
        raise HTTPException(status_code=400, detail="No content — upload a .docx/.txt or paste the questions.")
    content = content[:60000]

    if not parsed:
        parsed = _parse_assignments(content)
    if parsed:
        # Deterministic extraction of the prompts (+ rubric, if the doc had one) + AI translation only.
        items = _clean_assignments(await _translate_assignments(parsed))
    else:
        # Unstructured doc → let the AI find the prompts too.
        items = _clean_assignments(await _extract_items(ASSIGN_SYS, content, "assignments"))
    items = await _fill_assign_gaps(items)   # retry any language a batch missed
    if not items:
        raise HTTPException(status_code=422,
                            detail="No assignment prompts found — number each prompt (1., 2., …), use the "
                                   "Task Assigner template, or paste them one per line.")
    if replace:
        await db.execute(delete(AssignmentPrompt).where(AssignmentPrompt.video_id == video_id))
    order = 0 if replace else await _next_order(db, AssignmentPrompt, AssignmentPrompt.video_id, video_id)
    created = [AssignmentPrompt(video_id=video_id, question=it["question"], rubric=it["rubric"],
                               order_index=order + i)
               for i, it in enumerate(items)]
    for a in created:
        db.add(a)
    await db.commit()
    return {"added": len(created), "replaced": replace, "items": [_assign_dict(a) for a in created]}


@router.post("/sync-videos")
async def sync_videos(_: bool = Depends(require_admin), db: AsyncSession = Depends(get_db)):
    """Import known Cloudinary IDs (used by the learner site / WhatsApp) into the
    DB as per-language variants, matched by lesson title. Idempotent."""
    res = await db.execute(select(Video).options(selectinload(Video.language_variants)))
    videos = res.scalars().all()
    synced = []
    for v in videos:
        mapping = KNOWN_VIDEO_IDS.get(v.title)
        if not mapping:
            continue
        existing = {lv.language: lv for lv in v.language_variants}
        for lang, pid in mapping.items():
            if lang in existing:
                existing[lang].cloudinary_public_id = pid
            else:
                db.add(VideoLanguageVariant(video_id=v.id, language=lang, cloudinary_public_id=pid))
        if not v.cloudinary_public_id and mapping.get("en"):
            v.cloudinary_public_id = mapping["en"]
        synced.append({"title": v.title, "languages": sorted(mapping.keys())})

    # Import the existing quiz + assignment banks (deduped by English text).
    quizzes_added = assignments_added = 0
    legacy_path = Path(__file__).with_name("legacy_content.json")
    if legacy_path.exists():
        legacy = json.loads(legacy_path.read_text(encoding="utf-8"))
        by_title = {v.title: v for v in videos}

        for title, quizzes in legacy.get("quizzesByLesson", {}).items():
            v = by_title.get(title)
            if not v:
                continue
            ex = await db.execute(select(QuizQuestion).where(QuizQuestion.video_id == v.id))
            existing_q = {q.question.get("en") for q in ex.scalars().all()}
            base = await _next_order(db, QuizQuestion, QuizQuestion.video_id, v.id)
            for item in quizzes:
                if item["question"].get("en") in existing_q:
                    continue
                db.add(QuizQuestion(video_id=v.id, question=item["question"], options=item["options"],
                                    correct_index=item["correct_index"], order_index=base))
                base += 1
                quizzes_added += 1

        for title, prompts in legacy.get("assignmentsByLesson", {}).items():
            v = by_title.get(title)
            if not v:
                continue
            ex = await db.execute(select(AssignmentPrompt).where(AssignmentPrompt.video_id == v.id))
            existing_a = {a.question.get("en") for a in ex.scalars().all()}
            base = await _next_order(db, AssignmentPrompt, AssignmentPrompt.video_id, v.id)
            for item in prompts:
                if item["question"].get("en") in existing_a:
                    continue
                db.add(AssignmentPrompt(video_id=v.id, question=item["question"],
                                        rubric=item["rubric"], order_index=base))
                base += 1
                assignments_added += 1

    await db.commit()
    return {"videosSynced": len(synced), "quizzesAdded": quizzes_added,
            "assignmentsAdded": assignments_added}


async def _tree(course_id: str, db: AsyncSession) -> dict:
    return _course_tree(await _load_course(course_id, db))
